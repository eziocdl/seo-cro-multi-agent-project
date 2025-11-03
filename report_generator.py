

from pydantic import BaseModel, Field
from enum import Enum
from typing import List, Optional
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import black
from reportlab.lib.enums import TA_CENTER
from io import BytesIO 
import markdown
import os

# --- CONFIGURAÇÕES DE DADOS (Padrão Pydantic/Enum) ---

class ReportFormat(str, Enum):
    """Enum para formatos de relatório"""
    WORD = "word"
    PDF = "pdf"
    MARKDOWN = "markdown"

class ReportGenerationRequest(BaseModel):
    """Schema para a requisição da API de geração de relatório."""
    url: str
    report_format: ReportFormat = Field(
        default=ReportFormat.PDF, 
        description="Formato desejado para o relatório"
    )
    language: str = Field(
        default="pt-BR", 
        description="Idioma do relatório"
    )

# --- CLASSE DE ESTILO PARA PDF (ReportLab) ---

class ReportStyleConfig:
    """Configurações de estilo para geração de relatórios PDF (ReportLab)."""
    
    @staticmethod
    def get_pdf_styles():
        """Configura estilos personalizados para PDF."""
        styles = getSampleStyleSheet()
        
        styles.add(ParagraphStyle(
            name='H1Style',
            parent=styles['Title'],
            fontName='Helvetica-Bold',
            fontSize=18,
            textColor=black,
            alignment=TA_CENTER,
            spaceAfter=18
        ))
        
        styles.add(ParagraphStyle(
            name='H2Style',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            textColor=black,
            spaceAfter=6
        ))
        
        styles.add(ParagraphStyle(
            name='NormalStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            textColor=black,
            leading=12,
            spaceAfter=8
        ))
        
        return styles

# --- CLASSE PRINCIPAL DE GERAÇÃO DE ARQUIVOS (SRP) ---

class ReportGenerator:
    """Responsável por converter conteúdo Markdown em arquivos binários."""

    @staticmethod
    def generate_report(content: str, format: ReportFormat, url: str) -> BytesIO:
        """
        Gera relatório no formato especificado, retornando um objeto BytesIO.
        """
        if format == ReportFormat.WORD:
            return ReportGenerator._generate_word_report(content)
        
        elif format == ReportFormat.PDF:
            return ReportGenerator._generate_pdf_report(content)
        
        elif format == ReportFormat.MARKDOWN:
            return BytesIO(content.encode('utf-8'))
        
        raise ValueError(f"Formato {format.value} não suportado.")
    
    @staticmethod
    def _generate_word_report(content: str) -> BytesIO:
        """Gera relatório em Word (.docx)."""
        doc = docx.Document()
        doc.add_heading('Relatório de Auditoria Estratégica Digital', 0)
        
        # Converte Markdown para parágrafos simples
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line:
                doc.add_paragraph(line)
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    
    @staticmethod
    def _generate_pdf_report(content: str) -> BytesIO:
        """Gera relatório em PDF (ReportLab)."""
        file_stream = BytesIO()
        doc = SimpleDocTemplate(file_stream, pagesize=letter)
        styles = ReportStyleConfig.get_pdf_styles()
        elements = []
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                elements.append(Paragraph(line.replace('# ', ''), styles['H1Style']))
            elif line.startswith('## '):
                elements.append(Paragraph(line.replace('## ', ''), styles['H2Style']))
            elif line.startswith('- '):
                elements.append(Paragraph(f'&bull; {line.replace("- ", "").strip()}', styles['NormalStyle']))
            elif line:
                elements.append(Paragraph(line, styles['NormalStyle']))
            
            elements.append(Spacer(1, 6))
            
        doc.build(elements)
        file_stream.seek(0)
        return file_stream